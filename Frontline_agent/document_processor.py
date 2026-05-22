"""
Document Processing Service for Frontline Agent
Handles document parsing and text extraction from various file formats

Supports extraction of 100+ page PDFs. Uses pdfplumber (preferred) or PyPDF2 as fallback.
"""
import os
import hashlib
import logging
import re
from pathlib import Path
from typing import Dict, Optional, Tuple

from django.utils.text import get_valid_filename

logger = logging.getLogger(__name__)


# Magic-byte signatures per format. Each entry: (prefix_bytes, description).
# A file is accepted for a format only if its first bytes match one of the signatures.
_MAGIC_SIGNATURES = {
    'pdf': [(b'%PDF-', 'PDF')],
    'docx': [(b'PK\x03\x04', 'ZIP/OOXML')],  # DOCX is a ZIP; content-type check below confirms OOXML
    'doc': [(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1', 'OLE compound (legacy .doc)')],
    'html': [(b'<!DOCTYPE', 'HTML doctype'), (b'<html', 'HTML tag'),
             (b'<HTML', 'HTML tag'), (b'<?xml', 'XHTML')],
    # txt and md have no magic bytes — validated by UTF-8 decode heuristic below.
}


def _looks_like_text(data: bytes) -> bool:
    """Heuristic: decodable as UTF-8 or latin-1 AND contains few control characters."""
    try:
        sample = data[:4096]
        try:
            text = sample.decode('utf-8')
        except UnicodeDecodeError:
            text = sample.decode('latin-1')
        # Reject if >5% of characters are non-printable control bytes (excluding tab/newline/CR)
        bad = sum(1 for c in text if ord(c) < 32 and c not in '\t\n\r')
        return len(text) == 0 or (bad / max(len(text), 1)) < 0.05
    except Exception:
        return False


class DocumentProcessor:
    """Service for processing and extracting text from documents"""
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.txt': 'txt',
        '.md': 'md',
        '.html': 'html',
        '.htm': 'html',
    }
    
    # Max file size (50MB) - increased to support large documents with 100+ pages
    MAX_FILE_SIZE = 50 * 1024 * 1024
    
    @staticmethod
    def get_file_hash(file_path: str) -> str:
        """Calculate SHA256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    @staticmethod
    def get_file_format(filename: str) -> str:
        """Get file format from filename"""
        ext = Path(filename).suffix.lower()
        return DocumentProcessor.SUPPORTED_EXTENSIONS.get(ext, 'other')
    
    @staticmethod
    def validate_file(file_path: str, file_size: int) -> Tuple[bool, Optional[str]]:
        """Validate file before processing"""
        if file_size > DocumentProcessor.MAX_FILE_SIZE:
            return False, f"File size exceeds maximum allowed size ({DocumentProcessor.MAX_FILE_SIZE / 1024 / 1024}MB)"

        if not os.path.exists(file_path):
            return False, "File does not exist"

        return True, None

    @staticmethod
    def sanitize_filename(name: str) -> str:
        """Strip path separators and reduce to a safe filename (no path traversal)."""
        # Take basename only — defeats ../../etc/passwd style inputs
        base = os.path.basename(name or '')
        # Django helper: replaces spaces and strips anything not safe
        base = get_valid_filename(base) if base else ''
        # Collapse repeated dots / strip leading dots to avoid hidden files
        base = re.sub(r'\.{2,}', '.', base).lstrip('.')
        if not base:
            base = 'upload'
        # Hard cap length to keep FS happy
        if len(base) > 128:
            stem, dot, ext = base.rpartition('.')
            base = (stem[:120] + dot + ext) if dot else base[:128]
        return base

    @staticmethod
    def validate_content(data: bytes, filename: str) -> Tuple[bool, str, Optional[str]]:
        """
        Magic-byte / content validation. Returns (ok, detected_format, error).
        Rejects when the claimed extension and actual file content disagree.
        """
        ext_format = DocumentProcessor.get_file_format(filename)
        if ext_format == 'other' or ext_format not in DocumentProcessor.SUPPORTED_EXTENSIONS.values():
            return False, ext_format, f"Unsupported file type: {Path(filename).suffix}"

        if ext_format in ('txt', 'md'):
            if not _looks_like_text(data):
                return False, ext_format, "File does not look like valid text (binary or high-entropy content)"
            return True, ext_format, None

        signatures = _MAGIC_SIGNATURES.get(ext_format, [])
        for prefix, _desc in signatures:
            if data.startswith(prefix):
                # DOCX: confirm it's OOXML not a random ZIP by looking for the Content_Types entry
                if ext_format == 'docx' and b'[Content_Types].xml' not in data[:8192]:
                    continue
                return True, ext_format, None
        return False, ext_format, (
            f"File content does not match its extension ({ext_format}). "
            "Upload rejected to prevent disguised files."
        )
    
    @staticmethod
    def extract_text(file_path: str, file_format: str) -> Tuple[bool, str, Optional[str]]:
        """
        Extract text from document based on file format.
        
        Returns:
            Tuple of (success, extracted_text, error_message)
        """
        try:
            if file_format == 'txt':
                return DocumentProcessor._extract_txt(file_path)
            elif file_format == 'md':
                return DocumentProcessor._extract_txt(file_path)  # Markdown is text
            elif file_format == 'html':
                return DocumentProcessor._extract_html(file_path)
            elif file_format == 'pdf':
                return DocumentProcessor._extract_pdf(file_path)
            elif file_format == 'docx':
                return DocumentProcessor._extract_docx(file_path)
            elif file_format == 'doc':
                return DocumentProcessor._extract_doc(file_path)
            else:
                return False, '', f"Unsupported file format: {file_format}"
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}", exc_info=True)
            return False, '', str(e)
    
    @staticmethod
    def _extract_txt(file_path: str) -> Tuple[bool, str, Optional[str]]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return True, content, None
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                return True, content, None
            except Exception as e:
                return False, '', f"Failed to read text file: {str(e)}"
        except Exception as e:
            return False, '', f"Error reading text file: {str(e)}"
    
    @staticmethod
    def _extract_html(file_path: str) -> Tuple[bool, str, Optional[str]]:
        """Extract text from HTML file"""
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
            return True, text, None
        except ImportError:
            # Fallback to basic extraction if BeautifulSoup not available
            return DocumentProcessor._extract_txt(file_path)
        except Exception as e:
            return False, '', f"Error extracting HTML: {str(e)}"
    
    @staticmethod
    def _extract_pdf(file_path: str) -> Tuple[bool, str, Optional[str]]:
        """Extract text from PDF file - extracts ALL pages (up to 100+ pages).

        Falls back to OCR (Tesseract via pdf2image+pytesseract) when the text
        layer is empty or suspiciously thin — that's the symptom of a scanned
        / image-only PDF. Before OCR was wired, those PDFs got indexed as zero-
        text docs with no warning, breaking retrieval silently.
        """
        total_pages = 0
        # Try pdfplumber first (better for large/complex PDFs)
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"Using pdfplumber to extract text from PDF with {total_pages} pages")

                # Explicitly iterate through ALL pages - no limits
                for page_num in range(1, total_pages + 1):
                    try:
                        page = pdf.pages[page_num - 1]  # 0-indexed
                        page_text = page.extract_text()
                        if page_text:
                            # Embed a page marker so the chunker can later track
                            # which page each chunk belongs to. The marker uses
                            # form-feed (\x0c) which is rare in real text; the
                            # chunker strips it before storage.
                            text_parts.append(f"\x0c__PAGE_{page_num}__\n{page_text}")

                        # Log progress every 10 pages and at milestones
                        if page_num % 10 == 0 or page_num == total_pages:
                            logger.info(f"Extracted {page_num}/{total_pages} pages ({len(''.join(text_parts))} chars so far)")
                    except Exception as page_error:
                        logger.warning(f"Error extracting page {page_num}/{total_pages}: {page_error}, continuing...")
                        # Continue with other pages even if one fails
                        continue

                full_text = "\n".join(text_parts)
                logger.info(f"PDF extraction complete (pdfplumber): {len(full_text)} characters from {len(text_parts)}/{total_pages} pages")

                if len(text_parts) < total_pages:
                    logger.warning(f"Only extracted {len(text_parts)} out of {total_pages} pages. Some pages may have failed.")

                final_text, source = DocumentProcessor._maybe_ocr_pdf_fallback(
                    file_path, full_text, total_pages,
                )
                return True, final_text.strip(), None if source == 'text' else f'ocr_used={source}'
        except ImportError:
            logger.info("pdfplumber not available, trying PyPDF2...")
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2...")

        # Fallback to PyPDF2
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                total_pages = len(pdf_reader.pages)
                logger.info(f"Using PyPDF2 to extract text from PDF with {total_pages} pages")

                # Explicitly iterate through ALL pages - no limits
                for page_num in range(1, total_pages + 1):
                    try:
                        page = pdf_reader.pages[page_num - 1]  # 0-indexed
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)

                        # Log progress every 10 pages and at milestones
                        if page_num % 10 == 0 or page_num == total_pages:
                            logger.info(f"Extracted {page_num}/{total_pages} pages ({len(''.join(text_parts))} chars so far)")
                    except Exception as page_error:
                        logger.warning(f"Error extracting page {page_num}/{total_pages}: {page_error}, continuing...")
                        # Continue with other pages even if one fails
                        continue

                full_text = "\n".join(text_parts)
                logger.info(f"PDF extraction complete (PyPDF2): {len(full_text)} characters from {len(text_parts)}/{total_pages} pages")

                if len(text_parts) < total_pages:
                    logger.warning(f"Only extracted {len(text_parts)} out of {total_pages} pages. Some pages may have failed.")

                final_text, source = DocumentProcessor._maybe_ocr_pdf_fallback(
                    file_path, full_text, total_pages,
                )
                return True, final_text.strip(), None if source == 'text' else f'ocr_used={source}'
        except ImportError:
            return False, '', "Neither pdfplumber nor PyPDF2 library installed. Install with: pip install pdfplumber PyPDF2"
        except Exception as e:
            logger.error(f"Error extracting PDF with PyPDF2: {e}", exc_info=True)
            return False, '', f"Error extracting PDF: {str(e)}"

    @staticmethod
    def _maybe_ocr_pdf_fallback(file_path: str, current_text: str,
                                total_pages: int) -> Tuple[str, str]:
        """If the text layer is thin or empty, attempt OCR and return whichever
        result is longer. Returns ``(text, source)`` where source is
        ``'text'`` (no OCR ran or OCR didn't help), ``'ocr'`` (OCR text wins),
        or ``'unavailable'`` (OCR was needed but libraries aren't installed).

        Heuristic: OCR is attempted when the average page yields fewer than 30
        characters of text. That's far below any real document and the standard
        symptom of a scanned PDF. Skipped entirely when ``total_pages == 0``
        (parser failure — OCR won't help) or when the deployment opts out via
        ``settings.FRONTLINE_PDF_OCR_ENABLED = False``.
        """
        try:
            from django.conf import settings as _settings
            if not getattr(_settings, 'FRONTLINE_PDF_OCR_ENABLED', True):
                return current_text, 'text'
        except Exception:
            pass

        cur_len = len(current_text or '')
        threshold = max(100, int(total_pages) * 30)
        if total_pages <= 0 or cur_len >= threshold:
            return current_text, 'text'

        logger.info(
            "PDF text layer is thin (%d chars / %d pages, threshold %d) — attempting OCR fallback",
            cur_len, total_pages, threshold,
        )
        try:
            from pdf2image import convert_from_path  # type: ignore
            import pytesseract  # type: ignore
        except ImportError:
            logger.warning(
                "OCR fallback not available — install `pdf2image` + `pytesseract` "
                "(and Tesseract + Poppler at OS level) to enable scanned-PDF support. "
                "Returning %d chars of best-effort text-layer extraction.",
                cur_len,
            )
            return current_text, 'unavailable'

        try:
            ocr_parts: list[str] = []
            images = convert_from_path(file_path, dpi=200)
            for idx, image in enumerate(images, start=1):
                try:
                    page_text = pytesseract.image_to_string(image) or ''
                    if page_text.strip():
                        ocr_parts.append(page_text)
                    if idx % 5 == 0 or idx == len(images):
                        logger.info("OCR progress: %d/%d pages (%d chars so far)",
                                    idx, len(images), sum(len(p) for p in ocr_parts))
                except Exception as page_err:
                    logger.warning("OCR failed on page %d: %s — skipping", idx, page_err)
                    continue
            ocr_text = '\n'.join(ocr_parts).strip()
        except Exception as exc:
            logger.warning("OCR fallback failed entirely: %s — keeping text-layer extraction", exc)
            return current_text, 'text'

        # Only swap to OCR text if it's substantially better than what we had.
        if len(ocr_text) > max(cur_len * 2, threshold):
            logger.info("OCR fallback won: %d chars (text-layer was %d) — using OCR result",
                        len(ocr_text), cur_len)
            return ocr_text, 'ocr'
        logger.info("OCR ran but didn't beat the text layer (%d vs %d) — keeping original",
                    len(ocr_text), cur_len)
        return current_text, 'text'
    
    @staticmethod
    def _extract_docx(file_path: str) -> Tuple[bool, str, Optional[str]]:
        """Extract text from DOCX file - extracts ALL content including tables"""
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables (important for structured documents)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            logger.info(f"DOCX extraction complete: {len(full_text)} characters")
            return True, full_text, None
        except ImportError:
            return False, '', "python-docx library not installed. Install with: pip install python-docx"
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}", exc_info=True)
            return False, '', f"Error extracting DOCX: {str(e)}"
    
    @staticmethod
    def _extract_doc(file_path: str) -> Tuple[bool, str, Optional[str]]:
        """Extract text from DOC file (legacy Word format)"""
        # DOC files are more complex, may require additional libraries
        # For now, return error suggesting conversion to DOCX
        return False, '', "DOC format not directly supported. Please convert to DOCX format."
    
    @staticmethod
    def process_document(file_path: str, filename: str) -> Dict:
        """
        Process a document: validate, extract text, and return metadata.
        
        Returns:
            Dictionary with processing results
        """
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        
        # Validate file
        is_valid, error_msg = DocumentProcessor.validate_file(file_path, file_size)
        if not is_valid:
            return {
                'success': False,
                'error': error_msg,
                'file_hash': None,
                'file_format': None,
                'extracted_text': '',
            }
        
        # Get file format
        file_format = DocumentProcessor.get_file_format(filename)
        
        # Calculate file hash
        file_hash = DocumentProcessor.get_file_hash(file_path)
        
        # Extract text
        success, extracted_text, extract_error = DocumentProcessor.extract_text(file_path, file_format)
        
        if not success:
            return {
                'success': False,
                'error': extract_error or 'Failed to extract text',
                'file_hash': file_hash,
                'file_format': file_format,
                'extracted_text': '',
            }
        
        return {
            'success': True,
            'file_hash': file_hash,
            'file_format': file_format,
            'extracted_text': extracted_text,
            'file_size': file_size,
            'error': None,
        }





